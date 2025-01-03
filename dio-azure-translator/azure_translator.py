# -*- coding: utf-8 -*-
"""azure_translator.ipynb

Automatically generated by Colab.

Original file is located at
    https://colab.research.google.com/drive/1EshP3l4KVc2C4231DWi4yTwDUgV-UHBT

# Demonstração de tradução de texto usando o Azure

Este tutorial apresenta um exemplod de como traduzir um texto usando o *Azure Translator* e o *Azure OpenAI*.

## Instalando os pacotes necessários.
"""

# Commented out IPython magic to ensure Python compatibility.
# %%capture
# !pip install requests python-docx
# !pip install beautifulsoup4 openai langchain-openai

"""## Azure Translator Service

### Definindo a configuração do Azure Translator service

Preencha os parâmetros abaixo com a configuração do seu serviço *Azure Translator*. É necessário criar o serviço anteriormente no site do [Azure](https://azure.microsoft.com) e copiar a chave.
"""

import requests
import os
from docx import Document

# Insert here an API KEY for the Azure Translator service.
subscription_key = "AZURE_TRANSLATOR_KEY"
endpoint = "https://api.cognitive.microsofttranslator.com"
location = "eastus"
target_language ="pt-br"

def translate_text(text, target_language = 'pt-br'):
  path = '/translate'
  constructed_url = endpoint + path

  headers = {
      'Ocp-Apim-Subscription-Key': subscription_key,
      'Ocp-Apim-Subscription-Region': location,
      'Content-type': 'application/json',
      'X-ClientTraceId': str(os.urandom(16))
  }

  body = [{
      'text': text
  }]

  params = {
      'api-version': '3.0',
      'from': 'en',
      'to': target_language
  }

  request = requests.post(constructed_url, params = params, headers=headers, json=body)
  response = request.json()
  return response[0]["translations"][0]["text"]

"""### Simples teste do tradutor"""

translate_text("A soul in tension, learning to fly.")

"""### Tradução de um documento.

A função abaixo recebe o endereço de um documento em formado Word. Carregue um documento em seu ambiente do Colab para executar o exemplo.
"""

def translate_document(path):
  document = Document(path)
  full_text = []
  for paragraph in document.paragraphs:
    translated_text = translate_text(paragraph.text)
    full_text.append(translated_text)

    translated_doc = Document()
    for line in full_text:
      translated_doc.add_paragraph(line)
      path_translated = path.replace(".docx", f"-{target_language}.docx")
    translated_doc.save(path_translated)
  return full_text

translate_document("/content/learning-fly.docx")

"""## Azure OpenAI Service

### Carregando o conteúdo de um site.

Esta função carrega o conteudo de uma página web e faz uma limpeza removendo os comandos. O retorno é o texto da página.
"""

import requests
from bs4 import BeautifulSoup

def extract_text_from_url(url):
  response = requests.get(url)

  if response.status_code == 200:
    soup = BeautifulSoup(response.text, 'html.parser')
    for script_or_style in soup(["script", "style"]):
      script_or_style.decompose()

    text = soup.get_text(separator=' ')
    lines = (line.strip() for line in text.splitlines())
    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
    text = '\n'.join(chunk for chunk in chunks if chunk)

    return text
  else:
    print(f"Failed to retrieve the page. Status code: {response.status_code}")
    return None

extract_text_from_url("https://www.nasa.gov")

"""### Configurando o Langchain para acessar o Azure OpenAI Service.

A função abaixo faz a chamada do *Azure OpenAI Service*. Para utilizar é preciso provisionar um serviço no portal Azure e preencher as informações faltantes.
"""

from langchain_openai.chat_models.azure import AzureChatOpenAI

def translate_article(text, target_language = 'português'):
  client = AzureChatOpenAI(
      # Insert your Azure OpenAI service end-point.
      azure_endpoint = 'https://END-POINT.openai.azure.com/',
      # Insert your Azure OpenAI Service key.
      api_key = 'AZURE OPENAI SERVICE KEY',
      api_version = '2024-02-15-preview',
      deployment_name = 'gpt-4o-mini',
      max_retries= 0
  )

  messages = [
      ("system", "Você atua como um  tradutor de textos."),
      ("user", f"Traduza o {text} para o idioma {target_language} e responda em markdown.")
  ]

  client.invoke(messages)

  response = client.invoke(messages)
  print(response.content)
  return response.content

translate_article("A soul in tension, learning to fly.")

"""### Testando com um site"""

url = "https://www.nasa.gov"
text = extract_text_from_url(url)
article = translate_article(text)

print(article)